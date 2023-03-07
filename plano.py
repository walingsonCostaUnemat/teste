import streamlit as st
from docx import Document
from docx.shared import Inches

st.title('Gerador de Plano de Aula')
data = st.date_input('Data de início')
title = st.text_input('Título')
disciplina = st.text_input('Disciplina')
duracao = st.text_input('Duração')
serie = st.text_input('Série')
turma = st.text_input('Turma')

temas = st.text_area('Temas')
objetivos = st.text_area('Objetivos')
metodologia = st.text_area('Metodologia')
recursos = st.text_area('Recursos')
cronograma = st.text_area('Cronograma')
avaliacao = st.text_area('Avaliação')
observacoes = st.text_area('Observações')

submit_button = st.button('Gerar Plano de Aula')


def generate_docx(title, disciplina, duracao, serie, turma, temas, objetivos, metodologia, cronograma, avaliacao, observacoes):
    document = Document()
    document.add_heading(title, 0)
    document.add_heading('Escola Estadual São vicente de Paula', level=0)
    document.add_heading('Plano de aula')
    document.add_heading('Informações Gerais', level=1)
    document.add_paragraph(f'Data início: {data}')
    document.add_paragraph(f'Disciplina: {disciplina}')
    document.add_paragraph(f'Duração: {duracao}')
    document.add_paragraph(f'Série: {serie}')
    document.add_paragraph(f'Turma: {turma}')
    
    document.add_heading('Temas', level=1)
    for tema in temas.split('\n'):
        document.add_paragraph(tema)
    
    document.add_heading('Objetivos', level=1)
    for objetivo in objetivos.split('\n'):
        document.add_paragraph(objetivo)
        
    document.add_heading('Metodologia', level=1)
    for metod in metodologia.split('\n'):
        document.add_paragraph(metod)

    document.add_heading('Cronograma', level=1)
    for metod in cronograma.split('\n'):
        document.add_paragraph(metod)
    
    document.add_heading('Recursos', level=1)
    for metod in recursos.split('\n'):
        document.add_paragraph(metod)
        
    document.add_heading('Avaliação', level=1)
    for metod in avaliacao.split('\n'):
        document.add_paragraph(metod)
        
    document.add_heading('Observações', level=1)
    for metod in observacoes.split('\n'):
        document.add_paragraph(metod)
        
    document.save('plano_de_aula.docx')
if submit_button:
    generate_docx(title, disciplina, duracao, serie, turma, temas, objetivos, metodologia, cronograma, avaliacao, observacoes)
    with open('plano_de_aula.docx', 'rb') as f:
        data = f.read()
        st.download_button(label='Download', data=data, file_name='plano_de_aula.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
